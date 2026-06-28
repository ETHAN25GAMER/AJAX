from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def set_heading(paragraph, level, text, color=None):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x0a, 0x7d, 0x3c)
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x0a, 0x7d, 0x3c)
    elif level == 3:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    paragraph.paragraph_format.space_before = Pt(16)
    paragraph.paragraph_format.space_after = Pt(4)

def h1(doc, text):
    p = doc.add_paragraph()
    set_heading(p, 1, text)
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    set_heading(p, 2, text)
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    set_heading(p, 3, text)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "0a7d3c")
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "0a7d3c")
        tcPr.append(shd)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            row.cells[ci].text = val
            if ri % 2 == 1:
                tc = row.cells[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F2F8F4")
                tcPr.append(shd)
    doc.add_paragraph()

# ── Title block ──────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Ajax")
r.bold = True
r.font.size = Pt(32)
r.font.color.rgb = RGBColor(0x0a, 0x7d, 0x3c)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("WhatsApp AI Booking Assistant")
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

r3 = sub.add_run("\nWhat You Get & What We Need From You")
r3.font.size = Pt(12)
r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
sub.paragraph_format.space_after = Pt(20)

divider(doc)

# ── What Ajax Does ───────────────────────────────────────────────────────────
h1(doc, "What Ajax Does")
body(doc, "Ajax is an AI assistant that runs on your WhatsApp business number and handles customer conversations around the clock — no staff needed for routine enquiries. It books jobs, quotes prices, identifies pests from photos, and automatically nudges customers to re-book. Your team gets a web dashboard to manage everything, and technicians get a mobile app for the field.")

divider(doc)

# ── Part 1: WhatsApp Agent ───────────────────────────────────────────────────
h2(doc, "1.  WhatsApp AI Agent")
body(doc, "This is what your customers talk to. It handles the full booking lifecycle without any human involvement.")

h3(doc, "Booking")
bullet(doc, "Takes a customer from first message to confirmed appointment in one conversation")
bullet(doc, "Checks your real-time availability and proposes slots")
bullet(doc, "Books, reschedules, and cancels appointments")
bullet(doc, "Sends the customer a 6-character confirmation code")
bullet(doc, "Automatically assigns the job to your least-busy technician for that day")

h3(doc, "Pricing & Quotes")
bullet(doc, "Quotes prices from your exact price list — never makes up numbers")
bullet(doc, "Handles different tiers (e.g. one-off treatment vs. recurring plan)")
bullet(doc, "Flags which jobs need an on-site inspection before a firm price can be given")

h3(doc, "Pest Identification")
bullet(doc, "A customer can send a photo of a pest and Ajax will identify it, estimate severity, and suggest the right service")
bullet(doc, "Also works from a text description if no photo is available")

h3(doc, "Annual Maintenance Contracts (AMC)")
bullet(doc, "Checks whether a customer has an active AMC")
bullet(doc, "Takes renewal requests and flags them for your admin to confirm and collect payment")
bullet(doc, "Takes new subscription requests from customers who don't yet have a contract")

h3(doc, "Escalation")
bullet(doc, "Immediately flags urgent situations (bites/stings, structural risk, complaints) to a designated WhatsApp number")
bullet(doc, "Marks the conversation for your team to follow up — the customer is told a human will respond shortly")

h3(doc, "Automated Outreach")
body(doc, "Ajax sends proactive messages on your behalf via approved WhatsApp templates:")

add_table(doc,
    ["Message", "When it fires"],
    [
        ["Appointment reminder", "24 hours before every booked job"],
        ["Re-engagement nudge", "If a customer goes quiet mid-conversation"],
        ["En-route alert", "When a technician starts travel (includes a live tracking link)"],
        ["AMC renewal reminder", "X days before a contract renews (you set the lead time)"],
        ["AMC renewal follow-up", "7 days after the reminder if no response"],
        ["AMC upsell", "To past customers with no contract (max once per 90 days)"],
    ]
)

h3(doc, "Opt-Out Handling")
bullet(doc, "Customers can reply STOP at any time — Ajax immediately stops all promotional messages and confirms the opt-out")
bullet(doc, "Booking confirmations, reminders, and en-route alerts still go through (active service relationship)")
bullet(doc, "Customers can reply START to opt back in at any time")

divider(doc)

# ── Part 2: Admin Dashboard ──────────────────────────────────────────────────
h2(doc, "2.  Admin Dashboard (Web App)")
body(doc, "Your manager or owner logs in from any browser — desktop or mobile.")

add_table(doc,
    ["Section", "What you can do"],
    [
        ["Appointments", "See all upcoming jobs grouped by day; assign or reassign technicians; filter by status"],
        ["Dispatch", "Live view of every technician's GPS position, updated in real time as they travel"],
        ["Conversations", "Read the full WhatsApp message history for any customer"],
        ["Escalations", "Triage queue of conversations flagged for human attention; mark resolved when done"],
        ["Pricing", "Edit your price list directly — changes take effect on the next customer conversation"],
        ["AMC", "View and manage all annual maintenance contracts; see renewal status at a glance"],
        ["KPI Dashboard", "Revenue trends, job counts, technician performance, funnel metrics, AMC overview"],
        ["Users", "Invite admins and technicians; manage roles"],
        ["Settings", "Feature flags and system configuration"],
    ]
)

body(doc, "All changes sync in real time — if one admin assigns a technician, every other admin sees it immediately without refreshing.")

divider(doc)

# ── Part 3: Technician App ───────────────────────────────────────────────────
h2(doc, "3.  Technician Mobile App (PWA)")
body(doc, "Technicians install this on their phone like an app (no App Store required — it works in the browser). It shows only their own jobs.")

h3(doc, "What a technician can do on the app:")
bullet(doc, "See today's and tomorrow's assigned jobs in a clean list")
bullet(doc, "Tap a job to see the customer's name, address (with a direct Google Maps link), phone number (tap to call), pest type, and time slot")
bullet(doc, "Start travel — shares their live GPS with the customer via a link; the customer gets a WhatsApp message automatically; GPS updates every 30 seconds")
bullet(doc, "Mark job done or cancel it")
bullet(doc, "Add private tech notes (visible to admins, not customers)")
bullet(doc, "Upload before, after, and damage photos directly from their phone camera")
bullet(doc, "Flag an issue for dispatch (creates an escalation with urgency level)")

body(doc, "The customer's tracking page updates live and shows an estimated arrival time. It automatically deactivates once the job is marked complete.")

divider(doc)

# ── Privacy ──────────────────────────────────────────────────────────────────
h2(doc, "Privacy & Compliance")
bullet(doc, "Customer conversation history is automatically deleted after 6 months (configurable)")
bullet(doc, "Customers can opt out of promotional messages at any time by replying STOP")
bullet(doc, "A DPDP Act-compliant privacy notice is hosted on the app — customers can be referred to it")
bullet(doc, "All data is stored in a Supabase (Postgres) project in the India region")

divider(doc)

# ── What We Need ─────────────────────────────────────────────────────────────
h1(doc, "What We Need From You")
body(doc, "To set up Ajax for your business, we need the following. The more complete your answers, the faster we can go live.")

divider(doc)

h2(doc, "1.  Business Information")
bullet(doc, "Registered company name and brand name")
bullet(doc, "CIN / GSTIN (business registration number)")
bullet(doc, "Office address")
bullet(doc, "Areas / districts you cover")
bullet(doc, "Operating hours (days and times)")
bullet(doc, "Do you take bookings on public holidays?")

h2(doc, "2.  Your Price List  ★ most important — please be exact")
body(doc, "For every pest type you treat, we need:")
bullet(doc, "The pest name (e.g. cockroaches, rats, termites, bed bugs, mosquitoes…)")
bullet(doc, "Your price for each service tier:", "Standard — ")
bullet(doc, "Recurring / maintenance plan", "Plus — ")
bullet(doc, "Complex jobs requiring inspection (if applicable)", "Specialist — ")
bullet(doc, "Whether your prices are inclusive or exclusive of GST, and the GST rate")
bullet(doc, "Which pest types require an on-site inspection before you can quote")
bullet(doc, "What a standard job includes — e.g. \"30-day free re-treatment guarantee\"")
bullet(doc, "Roughly how long each job takes (e.g. Standard = 60 min, Plus = 90 min)")

h2(doc, "3.  Policies")
bullet(doc, "Cancellation and reschedule policy (e.g. \"first reschedule free; cancellation within 24h incurs a visit fee\")")
bullet(doc, "How customers pay (e.g. UPI, NEFT, cash, or card on completion — do you take deposits?)")

h2(doc, "4.  Agent Persona")
bullet(doc, "What should the assistant be called? (e.g. \"Asha\", \"Max\")")
bullet(doc, "Tone — warm and casual, or more formal? Are emoji okay?")
bullet(doc, "Which languages should it handle? (English, Hindi, Marathi, Tamil, Telugu, Kannada…)")
bullet(doc, "2–3 example messages you would normally send a customer — so we can match your style")
bullet(doc, "Anything the assistant must never say or promise")

h2(doc, "5.  Escalation Contact")
bullet(doc, "Which WhatsApp number should receive urgent alerts when a customer needs a real person? (Must be WhatsApp-enabled)")
bullet(doc, "What response time should Ajax promise customers for high-urgency situations?")

h2(doc, "6.  Your Team")
body(doc, "For every person who will use the admin dashboard or technician app:")

add_table(doc,
    ["Name", "Email (for login)", "Mobile", "Role"],
    [
        ["e.g. Priya Sharma", "priya@example.com", "+91 98765 43210", "Admin"],
        ["e.g. Ravi Kumar", "ravi@example.com", "+91 91234 56789", "Technician"],
    ]
)

bullet(doc, "Admin — sees all appointments and conversations; can edit pricing and manage team accounts")
bullet(doc, "Technician — sees only their own assigned jobs on the mobile app")

h2(doc, "7.  WhatsApp / Meta Access")
bullet(doc, "Is your WhatsApp number on WhatsApp Business API, the WhatsApp Business app, or a personal number? (We can help migrate — flag it early as it adds lead time)")
bullet(doc, "Who has admin access to your Meta Business Manager? We will need temporary partner access to set up the webhook and submit message templates.")

h2(doc, "8.  Existing AMC Contracts (if applicable)")
body(doc, "If you already have customers on recurring contracts, provide a list with:")
bullet(doc, "Customer name and WhatsApp number")
bullet(doc, "Pest type covered")
bullet(doc, "Contract start date and next renewal date")
bullet(doc, "Monthly or annual price")
body(doc, "We will import these so Ajax can handle renewals and reminders from day one.")

h2(doc, "9.  Legal / Compliance")
bullet(doc, "Name and email of your Data Protection Officer (or the person who handles privacy requests)")
bullet(doc, "Do you already have a privacy policy? If so, please share it — otherwise we will use our standard DPDP Act notice")
bullet(doc, "Pest control licence / PCO registration number (if you have one)")

divider(doc)

# ── Timeline ─────────────────────────────────────────────────────────────────
h2(doc, "Timeline")

add_table(doc,
    ["Phase", "What happens", "Typical time"],
    [
        ["You return the information above", "We configure Ajax for your business", "Day 1–2"],
        ["WhatsApp template submission", "We submit 7 message templates to Meta for approval", "1–48 hours"],
        ["Setup & testing", "We test booking, pricing, escalation, and tracking end to end", "Day 2–3"],
        ["Staff onboarding", "We walk your admin and technicians through the dashboard and app", "1 hour call"],
        ["Go live", "Ajax is live on your WhatsApp number", "Day 3–5"],
    ]
)

body(doc, "The main variable is Meta's template approval time. We submit them on day one so they are rarely the bottleneck.")

divider(doc)

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("Questions? Reply to this message or WhatsApp us directly.")
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
r.font.size = Pt(10)

# ── Save ──────────────────────────────────────────────────────────────────────
out = r"c:\Users\pllim\OneDrive\Desktop\python\PESTLLM\docs\Ajax_Client_Brief.docx"
doc.save(out)
print(f"Saved: {out}")
