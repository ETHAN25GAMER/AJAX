"""Render docs/CODEBASE_OVERVIEW.md into a styled .docx for Google Drive.

Drive auto-converts .docx to a native Google Doc on upload, so this is the
cleanest path to a readable, formatted doc. Styling mirrors make_client_brief.py
(GreenShield green headings, banded tables, monospace code blocks).
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "docs" / "CODEBASE_OVERVIEW.md"
OUT = ROOT / "docs" / "CODEBASE_OVERVIEW.docx"

GREEN = RGBColor(0x0A, 0x7D, 0x3C)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x88, 0x88, 0x88)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)


def add_inline(paragraph, text):
    """Render **bold**, `code`, and [label](link) inline markup into runs."""
    # Reduce links to their label; the URL isn't useful in a printed doc.
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    for token in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        elif token.startswith("*") and token.endswith("*"):
            r = paragraph.add_run(token[1:-1])
            r.italic = True
        elif token.startswith("`") and token.endswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = "Consolas"
            r.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        else:
            paragraph.add_run(token)


def heading(text, level):
    p = doc.add_paragraph()
    add_inline(p, text)
    size = {1: Pt(20), 2: Pt(15)}.get(level, Pt(12.5))
    color = DARK if level >= 3 else GREEN
    for r in p.runs:
        r.bold = True
        r.font.size = size
        r.font.color.rgb = color
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)


def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", "4"), ("w:space", "1"), ("w:color", "0a7d3c")):
        bottom.set(qn(k), v)
    pBdr.append(bottom)
    pPr.append(pBdr)


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    for k, v in (("w:val", "clear"), ("w:color", "auto"), ("w:fill", fill)):
        shd.set(qn(k), v)
    tcPr.append(shd)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        add_inline(cell.paragraphs[0], h)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(cell, "0a7d3c")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            add_inline(cell.paragraphs[0], val)
            if ri % 2 == 1:
                shade(cell, "F2F8F4")
    doc.add_paragraph()


def code_block(lines):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    for k, v in (("w:val", "clear"), ("w:color", "auto"), ("w:fill", "F4F6F4")):
        shd.set(qn(k), v)
    pPr.append(shd)
    r = p.add_run("\n".join(lines))
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(8)


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


# ── Title block ───────────────────────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("PestLLM")
r.bold, r.font.size, r.font.color.rgb = True, Pt(30), GREEN
s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = s.add_run("Codebase Overview")
rs.font.size, rs.font.color.rgb = Pt(15), RGBColor(0x44, 0x44, 0x44)
s.paragraph_format.space_after = Pt(14)
divider()

# ── Parse the markdown ────────────────────────────────────────────────────────
lines = SRC.read_text(encoding="utf-8").splitlines()

# Buffer for consecutive prose lines (markdown joins hard-wrapped lines into one
# paragraph; a blank line or any block element flushes the buffer).
prose: list[str] = []


def flush_prose():
    if not prose:
        return
    p = doc.add_paragraph()
    add_inline(p, " ".join(prose))
    p.paragraph_format.space_after = Pt(6)
    prose.clear()


i = 0
while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # Fenced code block
    if stripped.startswith("```"):
        flush_prose()
        buf = []
        i += 1
        while i < len(lines) and not lines[i].strip().startswith("```"):
            buf.append(lines[i])
            i += 1
        code_block(buf)
        i += 1
        continue

    # Table (header row followed by a |---| separator)
    if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[i + 1].strip()):
        flush_prose()
        headers = split_row(stripped)
        i += 2
        rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            rows.append(split_row(lines[i].strip()))
            i += 1
        add_table(headers, rows)
        continue

    if stripped.startswith("# "):
        flush_prose()
        heading(stripped[2:], 1)
    elif stripped.startswith("## "):
        flush_prose()
        heading(stripped[3:], 2)
    elif stripped.startswith("### "):
        flush_prose()
        heading(stripped[4:], 3)
    elif stripped == "---":
        flush_prose()
        divider()
    elif stripped.startswith(">"):
        flush_prose()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(12)
        add_inline(p, stripped.lstrip("> ").strip())
        for run in p.runs:
            run.italic = True
            run.font.color.rgb = GREY
        p.paragraph_format.space_after = Pt(6)
    elif re.match(r"^[-*] ", stripped):
        flush_prose()
        p = doc.add_paragraph(style="List Bullet")
        add_inline(p, stripped[2:])
        p.paragraph_format.space_after = Pt(2)
    elif stripped == "":
        flush_prose()
    else:
        prose.append(stripped)
    i += 1

flush_prose()
doc.save(OUT)
print(f"Saved: {OUT}")
